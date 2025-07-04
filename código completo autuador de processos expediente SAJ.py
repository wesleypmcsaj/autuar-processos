
from datetime import datetime
import os
import time
from pathlib import Path
import shutil
import re
import fitz  # PyMuPDF
from PyPDF2 import PdfReader, PdfWriter
import PyPDF2
import pdfplumber
import json
import threading
import win32print
import win32api

import imapclient
import pyzmail
import pyperclip

from collections import Counter

from reportlab.pdfgen import canvas
from io import BytesIO
import gspread  # Para acessar o Google Planilhas
from oauth2client.service_account import ServiceAccountCredentials  # Para autentica√ß√£o com conta de servi√ßo
from openpyxl import load_workbook
import openpyxl  # Para manipular o Excel sem alterar a formata√ß√£o
from PyPDF2 import PdfMerger
from send2trash import send2trash  # Biblioteca para mover arquivos para a lixeira
import pandas as pd
from docxtpl import DocxTemplate
from docx import Document

from selenium import webdriver
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import Select, WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from reportlab.lib.colors import red






# ----------------------------------------------------------------
#C√≥digo -1 [EXTRA], Verifica√ß√£o e limpeza de planilha para caso de execu√ß√£o do c√≥digo p√≥s erro;
# ----------------------------------------------------------------

#Implementado em 04/07/2025



# Caminhos dos arquivos
planilha_path = "dados_para_autuar_processos.xlsx"
bloco_notas_path = "livro_de_registros_pa_procuradores_judiciais.txt"

# Abrir a planilha mantendo a formata√ß√£o
wb = load_workbook(planilha_path)
ws = wb.active

# Identificar √≠ndice da coluna "Livro saj"
header = [cell.value for cell in ws[1]]
coluna_livro_saj_idx = header.index("Livro saj") + 1  # 1-based index

# Coletar valores da coluna "Livro saj" (linha 2 em diante)
valores_coluna = []
for row in ws.iter_rows(min_row=2, min_col=coluna_livro_saj_idx, max_col=coluna_livro_saj_idx):
    cell_value = row[0].value
    if cell_value and isinstance(cell_value, str) and cell_value.strip():
        valores_coluna.append(cell_value.strip())

# Se houver valores para processar
if valores_coluna:
    # Contar quantas vezes cada valor aparece na planilha
    contador_planilha = Counter(valores_coluna)

    # Ler o conte√∫do do bloco de notas
    with open(bloco_notas_path, "r", encoding="latin-1") as f:
        linhas_txt = f.readlines()

    # Preparar nova lista de linhas atualizadas
    linhas_txt_atualizadas = linhas_txt.copy()

    for valor, qtd_planilha in contador_planilha.items():
        # Encontrar √≠ndices das linhas com o valor e status "(j√° utilizado)"
        indices_alvo = [
            i for i, linha in enumerate(linhas_txt)
            if linha.strip().startswith(f"{valor},(j√° utilizado)")
        ]

        # Selecionar os √∫ltimos 'qtd_planilha' √≠ndices para altera√ß√£o
        for idx in reversed(indices_alvo[-qtd_planilha:]):
            linhas_txt_atualizadas[idx] = linhas_txt_atualizadas[idx].replace(
                ",(j√° utilizado)", ",(utiliz√°vel)", 1
            )

    # Escrever de volta o bloco de notas com as altera√ß√µes
    with open(bloco_notas_path, "w", encoding="latin-1") as f:
        f.writelines(linhas_txt_atualizadas)

    # Apagar colunas de A at√© I (colunas 1 a 9), exceto cabe√ßalho
    for row in ws.iter_rows(min_row=2, max_col=9):
        for cell in row:
            cell.value = None

    # Salvar a planilha mantendo a formata√ß√£o
    wb.save(planilha_path)

    print("Erro detectado em execu√ß√£o anterior do programa, procedimentos nescessarios tomados para nova execu√ß√£o da automa√ß√£o.")
else:
    print("N√£o foram encontrados erros de execu√ß√£o anteriores, c√≥digo pronto para copiar os dados registrados do Google Planilhas.")


#Ap√≥s excluir os registros, exclui arquivos pdf'S das pastas da automa√ß√£o, caso existam.

# Pastas a serem verificadas
pastas = ["docs numerados", "docs PJs", "docs processados"]

for pasta in pastas:
    if os.path.exists(pasta):
        for arquivo in os.listdir(pasta):
            caminho_completo = os.path.join(pasta, arquivo)
            if os.path.isfile(caminho_completo):  # S√≥ envia arquivos (n√£o pastas)
                try:
                    send2trash(caminho_completo)
                except Exception as e:
                    print(f"Erro ao mover {caminho_completo} para a lixeira: {e}")

# ----------------------------------------------------------------
#C√≥digo 0, copiar dados do Google Planilhas
# ----------------------------------------------------------------




print("Copiando os dados dispon√≠veis da planilha Google, aguarde.")

# Configura√ß√£o da autentica√ß√£o
def autenticar_google():
    # Define o escopo para acessar o Google Drive e Google Sheets
    scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/spreadsheets",
             "https://www.googleapis.com/auth/drive.file", "https://www.googleapis.com/auth/drive"]

    # Carrega as credenciais do arquivo JSON da conta de servi√ßo
    creds = ServiceAccountCredentials.from_json_keyfile_name('chave.json', scope)
    cliente = gspread.authorize(creds)
    return cliente

# Fun√ß√£o principal para copiar dados do Google Planilhas para o Excel
def copiar_dados_para_excel():
    # Autentica e conecta ao Google Planilhas
    cliente = autenticar_google()

    # Abre a planilha pelo ID e seleciona a aba "CAPTACOES"
    planilha = cliente.open_by_key("16_zlC5bRdyGTqFcVFvRIBCYzP-fjoPN9i64tD5DGe5c")
    aba = planilha.worksheet("CAPTACOES")

    # Define os cabe√ßalhos esperados, caso algum esteja duplicado ou em branco
    expected_headers = ["ID processo", "Procurador", "Org√£o"]
    dados = aba.get_all_records(expected_headers=expected_headers)

    # Abre o arquivo Excel existente sem alterar a formata√ß√£o
    arquivo_excel = "dados_para_autuar_processos.xlsx"
    workbook = openpyxl.load_workbook(arquivo_excel)
    sheet = workbook["CAPTACOES"]  # Nome da aba onde os dados ser√£o adicionados

    # Limpa as linhas de dados antigas (sem apagar o cabe√ßalho)
    sheet.delete_rows(2, sheet.max_row)

    # Insere os dados a partir da segunda linha
    for i, linha in enumerate(dados, start=2):  # Come√ßa na linha 2 para manter o cabe√ßalho original
        sheet.cell(row=i, column=1, value=linha["ID processo"])
        sheet.cell(row=i, column=2, value=linha["Procurador"])
        sheet.cell(row=i, column=3, value=linha["Org√£o"])

    # Salva o arquivo Excel com os dados atualizados
    workbook.save(arquivo_excel)
    print("Dados copiados com sucesso do Google Planilhas para o Excel.")

# Executa a fun√ß√£o para copiar os dados
copiar_dados_para_excel()



# ----------------------------------------------------------------
#C√≥digo 1 inserir livro SAJ nas c√©lulas correspondentes
# ----------------------------------------------------------------


def gerar_registros():
    # Cria o arquivo de registros com o padr√£o necess√°rio
    with open("livro_de_registros_pa_procuradores_judiciais.txt", "w") as arquivo:
        for livro in range(19, 51):  # Come√ßa no Livro 19 e vai at√© o Livro 50
            for pagina in range(1, 101):  # Vai da p√°gina 1 at√© a 100 (verso inclu√≠do)
                # Formata cada linha como "L <livro> P <pagina>, (utiliz√°vel)" e duplica a entrada
                arquivo.write(f"L {livro} P {pagina},(utiliz√°vel)\n")
                arquivo.write(f"L {livro} P {pagina},(utiliz√°vel)\n")  # Duplica√ß√£o da p√°gina com "(utiliz√°vel)"
                arquivo.write(f"L {livro} P {pagina} (v),(utiliz√°vel)\n")
                arquivo.write(f"L {livro} P {pagina} (v),(utiliz√°vel)\n")  # Duplica√ß√£o da p√°gina com "(v),(utiliz√°vel)"
    print("Arquivo de registros criado com sucesso.")


def verificar_e_preencher_excel():
    # Verifica se o arquivo "livro_de_registros_pa_procuradores_judiciais.txt" existe; se n√£o, gera o arquivo com o padr√£o especificado
    if not os.path.exists("livro_de_registros_pa_procuradores_judiciais.txt"):
        print("Arquivo 'livro_de_registros_pa_procuradores_judiciais.txt' n√£o encontrado. Gerando arquivo...")
        gerar_registros()

    # Abre o arquivo de Excel e a aba relevante
    arquivo_excel = "dados_para_autuar_processos.xlsx"
    workbook = openpyxl.load_workbook(arquivo_excel)
    sheet = workbook["CAPTACOES"]

    # L√™ o arquivo de registros
    with open("livro_de_registros_pa_procuradores_judiciais.txt", "r") as arquivo:
        registros = arquivo.readlines()

    # Itera sobre as linhas da planilha do Excel para preencher "Livro saj" com registros dispon√≠veis
    for row in range(2, sheet.max_row + 1):
        numero_processo = sheet.cell(row=row, column=1).value
        procurador = sheet.cell(row=row, column=2).value

        # Verifica se "numero processo" e "Procurador" est√£o preenchidos
        if numero_processo and procurador:
            # Encontra o primeiro registro com "(utiliz√°vel)"
            registro_disponivel = None
            for i, linha in enumerate(registros):
                if "(utiliz√°vel)" in linha:
                    registro_disponivel = linha.split(",")[0]
                    registros[i] = linha.replace("(utiliz√°vel)", "(j√° utilizado)")
                    break

            # Se n√£o houver registros dispon√≠veis, gera mais 50 livros e reinicia o preenchimento
            if not registro_disponivel:
                print("Todos os registros foram utilizados. Gerando mais 50 livros.")
                with open("livro_de_registros_pa_procuradores_judiciais.txt", "a") as arquivo:
                    for livro in range(51, 101):
                        for pagina in range(1, 101):
                            arquivo.write(f"L {livro} P {pagina},(utiliz√°vel)\n")
                            arquivo.write(f"L {livro} P {pagina},(utiliz√°vel)\n")
                            arquivo.write(f"L {livro} P {pagina} (v),(utiliz√°vel)\n")
                            arquivo.write(f"L {livro} P {pagina} (v),(utiliz√°vel)\n")
                return verificar_e_preencher_excel()

            # Insere o registro dispon√≠vel na coluna "Livro saj" (sexta coluna)
            sheet.cell(row=row, column=6, value=registro_disponivel)

    # Atualiza o arquivo de registros
    with open("livro_de_registros_pa_procuradores_judiciais.txt", "w") as arquivo:
        arquivo.writelines(registros)

    # Salva o arquivo Excel com as altera√ß√µes
    workbook.save(arquivo_excel)
    print("Livros de registros dos Pa's da procuradoria inseridos com sucesso!")


# Gera o arquivo de registros na primeira execu√ß√£o, se necess√°rio
verificar_e_preencher_excel()




# ----------------------------------------------------------------
#C√≥digo 2 baixar arquivos do ESAJ e inicio de completar informa√ß√µes nescessarias;
# ----------------------------------------------------------------






# Fun√ß√£o para atualizar o WebDriver e garantir que est√° atualizado
def update_webdriver():
    print("Atualizando o WebDriver e completando os dados da tabela. Por favor, aguarde.")
    service = Service(ChromeDriverManager().install())
    return service

def realizar_login(driver):
    with open("login_esaj_TJSP.txt", "r") as file:
        login = file.readline().strip()
        senha = file.readline().strip()

    driver.get(
        "https://esaj.tjsp.jus.br/sajcas/login?service=https%3A%2F%2Fesaj.tjsp.jus.br%2Fesaj%2Fj_spring_cas_security_check"
    )
    driver.find_element(By.XPATH, '//*[@id="usernameForm"]').send_keys(login)
    driver.find_element(By.XPATH, '//*[@id="passwordForm"]').send_keys(senha)
    driver.find_element(By.XPATH, '//*[@id="pbEntrar"]').click()
    print("Login inicial feito, aguardando tela do token...")
    time.sleep(5)

def buscar_token(driver):
    email_usuario = 'protocolosexpedientesajpmc@gmail.com'
    senha_app = 'otby gbko eryn fgmz'

    imap = imapclient.IMAPClient('imap.gmail.com', ssl=True)
    imap.login(email_usuario, senha_app)
    imap.select_folder('INBOX')

    hoje = datetime.now().strftime('%d-%b-%Y')
    criterio_busca = ['FROM', 'esaj@tjsp.jus.br']

    print("Aguardando o e-mail com o token...")

    tentativas_sem_token = 0

    while True:
        imap.noop()

        mensagens = imap.search(criterio_busca)
        if mensagens:
            mensagens.sort(reverse=True)
            for uid in mensagens:
                mensagem = imap.fetch([uid], ['BODY[]'])
                email = pyzmail.PyzMessage.factory(mensagem[uid][b'BODY[]'])

                if email.text_part:
                    conteudo = email.text_part.get_payload().decode(email.text_part.charset)
                    resultado = re.search(r'(\d{6})', conteudo)

                    if resultado:
                        token = resultado.group(1)
                        print(f"Token encontrado: {token}")

                        pyperclip.copy(token)
                        imap.move([uid], '[Gmail]/Trash')
                        print("E-mail movido para a lixeira.")
                        return token

        tentativas_sem_token += 1
        print(f"Token n√£o encontrado - tentativa {tentativas_sem_token}")

        # Se passou 61 tentativas (122s), clicar para reenviar o token
        if tentativas_sem_token >= 61:
            print("Token n√£o encontrado ap√≥s 120s, clicando para reenviar o token...")
            try:
                driver.find_element(By.XPATH, '//*[@id="btnReceberToken"]').click()
                print("Cliquei para reenviar o token!")
            except:
                print("N√£o encontrei o bot√£o de reenviar token.")

            tentativas_sem_token = 0  # Resetar tentativas

        time.sleep(2)

def inserir_token(driver, token):
    campo_token = driver.find_element(By.XPATH, '//*[@id="tokenInformado"]')
    campo_token.clear()
    campo_token.send_keys(token)
    driver.find_element(By.XPATH, '//*[@id="btnEnviarToken"]').click()
    print("Token inserido, aguardando valida√ß√£o...")

    try:
        WebDriverWait(driver, 5).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="avisoTokenInvalido"]/span'))
        )
        print("Token inv√°lido detectado, buscando novo token...")
        return False  # Token inv√°lido
    except:
        print("Token aceito!")
        return True  # Token v√°lido

def login_esaj(driver):
    try:
        realizar_login(driver)

        while True:
            token = buscar_token(driver)
            if inserir_token(driver, token):
                break  # Sai do loop se o token for v√°lido
            else:
                print("Repetindo processo de busca de novo token...")

        print("Login realizado no sistema E-SAJ")
        time.sleep(5)

    except Exception as e:
        print(f"Erro durante o login: {e}")
        raise

# Fun√ß√£o para acessar o site de consulta
def acessar_site(driver):
    driver.get("https://esaj.tjsp.jus.br/cpopg/open.do")
    time.sleep(1)


# Fun√ß√£o para fechar janelas extras, mantendo apenas a janela principal aberta
def fechar_janelas_extras(driver, janela_principal):
    for handle in driver.window_handles:
        if handle != janela_principal:
            driver.switch_to.window(handle)
            driver.close()
    driver.switch_to.window(janela_principal)


# Caminho da pasta de downloads
downloads_path = Path.home() / "Downloads"


# Fun√ß√£o para mover o arquivo baixado para a pasta "docs PJs" e salvar o caminho em um bloco de notas
def mover_arquivo_downloads(destino_pasta):
    docs_pjs_path = Path(destino_pasta)

    # Verifica se a pasta "docs PJs" existe, se n√£o, cria
    docs_pjs_path.mkdir(parents=True, exist_ok=True)

    # Filtra apenas arquivos PDF com a data atual no diret√≥rio de downloads
    today = datetime.now().date()
    pdf_files_today = [
        f for f in downloads_path.glob('*.pdf')
        if datetime.fromtimestamp(f.stat().st_mtime).date() == today
    ]

    # Se houver arquivos PDF da data atual, move o mais recente para "docs PJs"
    if pdf_files_today:
        latest_file = max(pdf_files_today, key=os.path.getmtime)
        shutil.move(str(latest_file), str(docs_pjs_path / latest_file.name))

    # Salva o caminho da pasta "docs PJs" em um bloco de notas
    with open("diretorio_docs_pjs.txt", "w") as file:
        file.write(str(docs_pjs_path))


# Fun√ß√£o para consultar o processo e salvar as informa√ß√µes no Excel
def consultar_processo(driver, id_processo, row_index, sheet):
    if id_processo:
        match = re.match(r"(\d{7})-(\d{2})\.(\d{4})\.8\.26\.(\d{4})", id_processo)
        if match:
            numero_digito_ano = f"{match.group(1)}-{match.group(2)}.{match.group(3)}"
            foro_numero_unificado = match.group(4)

            driver.get("https://esaj.tjsp.jus.br/cpopg/open.do?gateway=true")
            time.sleep(1)

            # Seleciona a op√ß√£o "N√∫mero do Processo" no dropdown de pesquisa
            select_element = driver.find_element(By.XPATH, '//*[@id="cbPesquisa"]')
            select = Select(select_element)
            select.select_by_visible_text("N√∫mero do Processo")

            # Preenche os campos com os valores formatados
            driver.find_element(By.XPATH, '//*[@id="numeroDigitoAnoUnificado"]').send_keys(numero_digito_ano)
            driver.find_element(By.XPATH, '//*[@id="foroNumeroUnificado"]').send_keys(foro_numero_unificado)

            # Clica no bot√£o de consultar processo
            driver.find_element(By.XPATH, '//*[@id="botaoConsultarProcessos"]').click()
            time.sleep(10)

            # Classe do processo
            classe_texto = driver.find_element(By.XPATH, '//*[@id="classeProcesso"]').text
            sheet.cell(row=row_index, column=4, value=classe_texto)

            # Assunto do processo
            assunto_texto = driver.find_element(By.XPATH, '//*[@id="assuntoProcesso"]').text
            sheet.cell(row=row_index, column=5, value=assunto_texto)

            # Nome da parte
            nome_parte_completo = driver.find_element(By.XPATH,
                                                      '//*[@id="tablePartesPrincipais"]/tbody/tr[1]/td[2]').text
            if "Advogada:" in nome_parte_completo:
                nome_parte = nome_parte_completo.split("Advogada:")[0].strip()
            elif "Advogado:" in nome_parte_completo:
                nome_parte = nome_parte_completo.split("Advogado:")[0].strip()
            else:
                nome_parte = nome_parte_completo
            sheet.cell(row=row_index, column=7, value=nome_parte)

            # Salva o arquivo atualizado
            sheet.parent.save("dados_para_autuar_processos.xlsx")

            try:
                # Abre a pasta de documentos do processo
                driver.find_element(By.XPATH, '//*[@id="linkPasta"]').click()
                WebDriverWait(driver, 10).until(lambda d: len(d.window_handles) > 1)

                janela_principal = driver.current_window_handle

                # Alterna para a nova janela (pop-up)
                for handle in driver.window_handles:
                    if handle != driver.current_window_handle:
                        driver.switch_to.window(handle)
                        break

                # Clica no bot√£o "Selecionar todos os documentos"
                driver.find_element(By.XPATH, '//*[@id="selecionarButton"]').click()
                time.sleep(3)

                # Clica no bot√£o "Salvar"
                driver.find_element(By.XPATH, '//*[@id="salvarButton"]').click()
                time.sleep(3)

                # Aguarda at√© que o bot√£o "Continuar" esteja vis√≠vel e clic√°vel e clica nele
                continuar_button = WebDriverWait(driver, 20).until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="botaoContinuar"]'))
                )
                continuar_button.click()
                time.sleep(10)
                # Tempo equivalente a 1 hora
                # Aguarda at√© que o bot√£o de download esteja dispon√≠vel e clica nele
                download_button = WebDriverWait(driver, 3600).until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="btnDownloadDocumento"]'))
                )
                download_button.click()
                time.sleep(10)

                # Fecha a janela pop-up ap√≥s o download
                driver.close()
                driver.switch_to.window(janela_principal)

                # Espera at√© que o arquivo .crdownload seja removido
                while any(downloads_path.glob("*.crdownload")):
                    time.sleep(1)

                mover_arquivo_downloads("docs PJs")

            except Exception as e:
                print(f"Erro durante o processo de download: {e}")

        else:
            print(f"Aviso: ID do processo '{id_processo}' n√£o est√° no formato esperado.")
    else:
        print("Aviso: O ID do processo est√° vazio. Linha ignorada.")


# Fun√ß√£o principal para processar todos os dados
def processar_dados():
    wb = load_workbook("dados_para_autuar_processos.xlsx")
    sheet = wb["CAPTACOES"]

    chrome_options = Options()
    service = update_webdriver()
    driver = webdriver.Chrome(service=service, options=chrome_options)

    login_esaj(driver)

    for i, row in enumerate(sheet.iter_rows(min_row=2, max_col=8, values_only=True), start=2):
        id_processo, procurador, orgao = row[0], row[1], row[2]

        # Verifica√ß√£o se "ID processo" come√ßa com "///"
        if id_processo and id_processo.startswith("///"):
            # Extra√ß√£o e divis√£o das informa√ß√µes
            partes = id_processo[3:].strip().split(" / ")
            if len(partes) >= 4:
                numero_processo, classe, assunto, nome_parte = partes[:4]

                # Preenche as colunas espec√≠ficas
                sheet.cell(row=i, column=4, value=classe)
                sheet.cell(row=i, column=5, value=assunto)
                sheet.cell(row=i, column=7, value=nome_parte)

                # Preenche a 8¬™ coluna com a formata√ß√£o solicitada
                if procurador:
                    informacoes_completas = (
                        f"A√ß√£o envolvendo {nome_parte} e PREFEITURA MUNICIPAL DE CARAPICUIBA {classe} {assunto} {numero_processo}"
                    )
                else:
                    informacoes_completas = (
                        f"Solicita√ß√£o de provid√™ncias/ solicita√ß√£o de informa√ß√µes envolvendo {nome_parte} e PREFEITURA MUNICIPAL DE CARAPICUIBA {classe} {assunto} {numero_processo}"
                    )
                sheet.cell(row=i, column=8, value=informacoes_completas)

            # Pula o processamento no Selenium para essas linhas
            print(f"Linha {i} iniciada com '///' foi processada sem usar o Selenium.")
            continue

        # Apenas processa no Selenium se "ID processo" n√£o come√ßar com "///"
        if orgao == "TJSP" and id_processo:
            consultar_processo(driver, id_processo, i, sheet)

    # Salva as altera√ß√µes no arquivo Excel
    wb.save("dados_para_autuar_processos.xlsx")
    driver.quit()


# Chama a fun√ß√£o principal
if __name__ == "__main__":
    processar_dados()



# ----------------------------------------------------------------
#C√≥digo 3 - Finalizar as informa√ß√µes nescessarias para autuar os processos no GIAP
# ----------------------------------------------------------------







# Mensagem inicial
print("As informa√ß√µes necess√°rias para autuar os processos administrativos Judiciais est√£o sendo completadas, por favor aguarde.")

# Carregar o arquivo Excel
arquivo_excel = "dados_para_autuar_processos.xlsx"
workbook = openpyxl.load_workbook(arquivo_excel)
planilha = workbook.active

# Identificar os √≠ndices das colunas (usando os nomes fornecidos)
coluna_id_processo = 1  # "ID processo"
coluna_procurador = 2  # "Procurador"
coluna_orgao = 3  # "Org√£o"
coluna_classe = 4  # "Classe"
coluna_assunto = 5  # "Assunto"
coluna_livro_saj = 6  # "Livro saj"
coluna_nome_parte = 7  # "Nome da parte"
coluna_info_completa = 8  # "Informacoes completas para autuar o PA"

# Fun√ß√£o para extrair o n√∫mero do processo judicial
def extrair_numero_processo(id_processo):
    if id_processo is None:
        return ""
    id_processo = str(id_processo)
    if "///" in id_processo:
        partes = id_processo.split("/")
        for parte in partes:
            if parte and len(parte) >= 25 and "-" in parte and "." in parte:
                return parte
    elif len(id_processo) >= 25 and "-" in id_processo and "." in id_processo:
        return id_processo
    return ""

# Iterar pelas linhas (ignorar a primeira linha, que √© o cabe√ßalho)
for linha in range(2, planilha.max_row + 1):
    procurador = planilha.cell(row=linha, column=coluna_procurador).value
    orgao = planilha.cell(row=linha, column=coluna_orgao).value
    nome_parte = planilha.cell(row=linha, column=coluna_nome_parte).value
    classe = planilha.cell(row=linha, column=coluna_classe).value
    assunto = planilha.cell(row=linha, column=coluna_assunto).value
    id_processo = planilha.cell(row=linha, column=coluna_id_processo).value
    livro_saj = planilha.cell(row=linha, column=coluna_livro_saj).value

    numero_processo = extrair_numero_processo(id_processo)  # Extrair n√∫mero do processo judicial

    # Adicionar espa√ßo simples e valor da coluna "Livro saj" (ou apenas espa√ßo se vazio)
    complemento_livro_saj = f" {livro_saj}" if livro_saj else " "

    # Condicional para quando o "Org√£o" for "TJSP"
    if orgao == "TJSP":
        if procurador:  # Se o "Procurador" estiver preenchido
            texto = (f"A√ß√£o Judicial envolvendo {nome_parte} e PREFEITURA MUNICIPAL DE CARAPICUIBA. "
                     f"{classe}   {assunto}  - n√∫mero do processo: {numero_processo} {complemento_livro_saj} {procurador}")
        else:  # Se o "Procurador" estiver vazio
            texto = (f"Solicita√ß√£o de provid√™ncias/ Solicita√ß√£o de Informa√ß√µes envolvendo {nome_parte} "
                     f"e PREFEITURA MUNICIPAL DE CARAPICUIBA   {classe}   {assunto} - n√∫mero do processo: {numero_processo}")
    else:  # Quando o "Org√£o" for diferente de "TJSP"
        texto = f"Solicita√ß√£o de provid√™ncias/ Solicita√ß√£o de Informa√ß√µes envolvendo {id_processo} {complemento_livro_saj} {procurador}"

    # Inserir o texto na 8¬™ coluna ("Informacoes completas para autuar o PA")
    planilha.cell(row=linha, column=coluna_info_completa).value = texto

# Salvar o arquivo Excel atualizado
workbook.save(arquivo_excel)

# Mensagem final
print("Todas as informa√ß√µes necess√°rias para autuar os processos foram inseridas.")


print("Os documentos est√£o sendo divididos em suas respectivas peti√ß√µes e decis√µes. Esse processo tende a demorar.")



# ----------------------------------------------------------------
#C√≥digo 4 dividir os PDF's em suas respectivas peti√ß√µes e Decis√µes alterado em 24-06-2025
# ----------------------------------------------------------------







#Novo c√≥digo, ele possui agora uma varia√ß√£o  utilizando "pdfplumber" para caso o PyPDF n√£o consiga processar algum arquivo em especifico,
#como se fosse uma contramedida.


# Configura√ß√£o de pastas
PASTA_ORIGEM = "docs PJs"
PASTA_SAIDA = "docs processados"
os.makedirs(PASTA_SAIDA, exist_ok=True)

# Termos de busca aprimorados (combinando ambos c√≥digos)
TERMINOS_PETICAO = [
    "Nestes termos p. deferimento", "Termos em que, Pede Deferimento.", "Pede Deferimento",
    "Nestes termos, pede deferimento", "Nesses termos pede e aguarda deferimento",
    "Termos em que, Pede e Espera Deferimento.", "Termosemque, PedeeEsperaDeferimento.",
    "Termosemque, PedeeEsperaDeferimento", "Termos em que, Pede e Espera Deferimento",
    "Termos em que,", "p. deferimento", "Pede e Espera Deferimento.", "Termosemque,",
    "PedeeEsperaDeferimento.", "PedeeEsperaDeferimento", "Nestes termos, pede e espera deferimento",
    "Nestes termos, espera deferimento", "Nesses termos pede deferimento.", "Nesses termos, pede deferimento.",
    "Nesses termos pede deferimento", "Nesses termos, pede deferimento", "Nesses termos pede e aguarda deferimento",
    "Nesses termos, pede e aguarda deferimento", "Nesses termos pede e aguarda deferimento."
]

TERMO_DECISAO = "DECIS√ÉO"
TERMO_DESPACHO = "DESPACHO"
TERMO_OFICIO = "OF√çCIO"
TERMO_TRIBUNAL = "TRIBUNAL DE JUSTI√áA DO ESTADO DE S√ÉO PAULO"

TERMO_JUIZ = [
    "Juiz(a) de Direito: Dr(a)", "Juiz de Direito:", "Ju√≠za de Direito:",
    "Juiz(a) de Direito", "Juiz de Direito", "Ju√≠za de Direito"
]


def localizar_paginas(doc, termos, termos_adicionais=None, ultima_ocorrencia=False):
    """
    Localiza p√°ginas em um documento PDF com base em termos principais e adicionais.

    Args:
        doc: Documento PDF carregado com fitz.
        termos: Lista de termos principais a serem buscados.
        termos_adicionais: Lista de termos adicionais a serem buscados (opcional).
        ultima_ocorrencia: Se True, retorna apenas a √∫ltima ocorr√™ncia encontrada.

    Returns:
        Lista de √≠ndices das p√°ginas encontradas ou a √∫ltima p√°gina se ultima_ocorrencia for True.
    """
    paginas_encontradas = []
    for i, pagina in enumerate(doc):
        texto = pagina.get_text()

        if any(termo in texto for termo in termos):
            if termos_adicionais is None or any(termo in texto for termo in termos_adicionais):
                paginas_encontradas.append(i)

    return [paginas_encontradas[-1]] if ultima_ocorrencia and paginas_encontradas else paginas_encontradas


def processar_com_pymupdf(caminho_arquivo, nome_base):
    """Processa o PDF usando PyMuPDF (fitz) como m√©todo principal."""
    doc = fitz.open(caminho_arquivo)
    try:
        # Localizar p√°ginas da peti√ß√£o
        paginas_peticao = []
        for i, pagina in enumerate(doc):
            texto = pagina.get_text()
            if any(termo in texto for termo in TERMINOS_PETICAO):
                paginas_peticao = list(range(0, i + 1))
                break

        if not paginas_peticao:
            paginas_peticao = list(range(0, min(15, len(doc))))

        # Localizar decis√£o/despacho/of√≠cio
        paginas_decisao = localizar_paginas(doc, [TERMO_DECISAO], TERMO_JUIZ, ultima_ocorrencia=True)
        if not paginas_decisao:
            paginas_decisao = localizar_paginas(doc, [TERMO_DESPACHO], TERMO_JUIZ, ultima_ocorrencia=True)
        if not paginas_decisao:
            paginas_decisao = localizar_paginas(doc, [TERMO_OFICIO], ultima_ocorrencia=True)

        # Criar documentos separados
        doc_peticao = fitz.open()
        doc_decisao = fitz.open()

        for p in paginas_peticao:
            doc_peticao.insert_pdf(doc, from_page=p, to_page=p)
        for p in paginas_decisao:
            doc_decisao.insert_pdf(doc, from_page=p, to_page=p)

        # Combinar resultados
        doc_final = fitz.open()
        if len(doc_peticao) > 0:
            doc_final.insert_pdf(doc_peticao)
        if len(doc_decisao) > 0:
            doc_final.insert_pdf(doc_decisao)

        if len(doc_final) > 0:
            caminho_saida = os.path.join(PASTA_SAIDA, f"{nome_base}.pdf")
            doc_final.save(caminho_saida)
            print(f"Processado com PyMuPDF: {caminho_saida}")
            return True
        else:
            raise Exception("Nenhuma p√°gina v√°lida encontrada.")

    except Exception as e:
        print(f"Erro no processamento principal: {e}")
        return False
    finally:
        doc.close()
        if 'doc_peticao' in locals(): doc_peticao.close()
        if 'doc_decisao' in locals(): doc_decisao.close()
        if 'doc_final' in locals(): doc_final.close()


def processar_com_metodo_alternativo(caminho_arquivo, nome_base):
    """M√©todo alternativo usando PyPDF2 + pdfplumber quando o principal falha."""
    print(f"‚ö†Ô∏è Tentando m√©todo alternativo para: {nome_base}")

    try:
        paginas_peticao = []
        paginas_decisao = []

        with pdfplumber.open(caminho_arquivo) as pdf:
            for i, page in enumerate(pdf.pages):
                texto = page.extract_text() or ""

                if not paginas_peticao and any(t in texto for t in TERMINOS_PETICAO):
                    paginas_peticao = list(range(0, i + 1))

                if any(t in texto for t in [TERMO_DECISAO, TERMO_DESPACHO, TERMO_OFICIO]):
                    if any(j in texto for j in TERMO_JUIZ):
                        paginas_decisao = [i]  # Mant√©m apenas a √∫ltima

        if not paginas_peticao:
            paginas_peticao = list(range(0, min(15, len(pdf.pages))))

        reader = PyPDF2.PdfReader(caminho_arquivo)
        writer = PyPDF2.PdfWriter()

        for i in paginas_peticao:
            writer.add_page(reader.pages[i])
        for i in paginas_decisao:
            writer.add_page(reader.pages[i])

        if writer.pages:
            caminho_saida = os.path.join(PASTA_SAIDA, f"{nome_base}.pdf")
            with open(caminho_saida, "wb") as f_out:
                writer.write(f_out)
            print(f"‚úÖ Processado com m√©todo alternativo: {caminho_saida}")
            return True
        else:
            raise Exception("Nenhuma p√°gina v√°lida encontrada no m√©todo alternativo.")

    except Exception as e:
        print(f"Erro no m√©todo alternativo: {e}")
        return False


def mover_para_lixeira(caminho_arquivo):
    """Tenta mover para a lixeira, se falhar remove permanentemente."""
    try:
        # Tentativa de mover para lixeira (Windows)
        if os.name == 'nt':
            import send2trash
            send2trash.send2trash(caminho_arquivo)
            print(f"Enviado para lixeira: {caminho_arquivo}")
        else:
            # Em outros sistemas, remove permanentemente
            os.remove(caminho_arquivo)
            print(f"Removido permanentemente: {caminho_arquivo}")
    except Exception as e:
        print(f"Falha ao mover para lixeira, removendo permanentemente: {e}")
        os.remove(caminho_arquivo)


def main():
    falhas = []
    sucessos = 0

    for arquivo in os.listdir(PASTA_ORIGEM):
        if not arquivo.endswith(".pdf"):
            continue

        caminho_arquivo = os.path.join(PASTA_ORIGEM, arquivo)
        nome_base = Path(arquivo).stem

        try:
            # Tentar primeiro com PyMuPDF
            sucesso = processar_com_pymupdf(caminho_arquivo, nome_base)

            if not sucesso:
                # Se falhar, tentar m√©todo alternativo
                sucesso = processar_com_metodo_alternativo(caminho_arquivo, nome_base)
                if sucesso:
                    falhas.append((arquivo, "M√©todo alternativo"))
                else:
                    falhas.append((arquivo, "Falha total"))
            else:
                sucessos += 1

            # Mover arquivo original para lixeira/apagar
            mover_para_lixeira(caminho_arquivo)

        except Exception as e:
            print(f"Erro inesperado ao processar {arquivo}: {e}")
            falhas.append((arquivo, "Erro inesperado"))

    # Relat√≥rio final
    print("\n=== RELAT√ìRIO FINAL ===")
    print(f"‚úÖ {sucessos} arquivos processados com sucesso pelo m√©todo principal")

    if falhas:
        print("\n‚ö†Ô∏è Arquivos com problemas:")
        for arquivo, motivo in falhas:
            print(f" - {arquivo}: {motivo}")
    else:
        print("\nüéâ Todos os arquivos foram processados com sucesso!")


if __name__ == "__main__":
    main()





# ----------------------------------------------------------------
#C√≥digo 5 - numera√ß√£o das folhas dos arquivos .PDF
# ----------------------------------------------------------------


def carregar_dados_txt(caminho_txt):
    """L√™ a sigla e a matr√≠cula de um arquivo .txt"""
    with open(caminho_txt, "r") as f:
        linhas = f.readlines()
    sigla = linhas[0].strip()  # Primeira linha: Sigla
    matricula = linhas[1].strip()  # Segunda linha: Matr√≠cula
    return sigla, matricula

def ajustar_orientacao_e_numerar(input_path, output_path, sigla, matricula, inicio=2, fonte="Helvetica", tamanho_fonte=12):
    """Ajusta a orienta√ß√£o das p√°ginas para retrato e adiciona numera√ß√£o personalizada com l√≥gica de frente e verso."""
    reader = PdfReader(input_path)
    writer = PdfWriter()

    numero_pagina = inicio  # Inicia a numera√ß√£o pela p√°gina 2

    for indice, pagina in enumerate(reader.pages):
        # Verifica e ajusta a rota√ß√£o da p√°gina
        if pagina.get("/Rotate") in [90, 270]:  # P√°ginas na horizontal
            pagina.rotate(0)  # Ajusta para retrato

        # Obt√©m dimens√µes da p√°gina
        largura_pagina = float(pagina.mediabox[2])
        altura_pagina = float(pagina.mediabox[3])

        # Buffer para gerar a numera√ß√£o com ReportLab
        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=(largura_pagina, altura_pagina))

        # Define a numera√ß√£o da p√°gina
        if indice % 2 == 0:  # P√°ginas pares (frente)
            texto_numero = f"Folha {numero_pagina}"
        else:  # P√°ginas √≠mpares (verso)
            texto_numero = f"Folha {numero_pagina} verso"
            numero_pagina += 1  # Incrementa a numera√ß√£o somente ap√≥s o verso

        texto_sigla = f"{sigla}"
        texto_matricula = f"{matricula}"

        # Defini√ß√µes do quadro branco
        margem_direita = 60  # Dist√¢ncia da borda direita
        margem_topo = 30  # Dist√¢ncia do topo
        largura_quadro = 80  # Largura do quadro branco
        altura_quadro = 30  # Altura do quadro branco
        margem_interna = 10  # Margem interna para o texto dentro do quadro

        # Calcula a posi√ß√£o do quadro branco no canto superior direito
        pos_x_quadro = largura_pagina - margem_direita - largura_quadro
        pos_y_quadro = altura_pagina - margem_topo - altura_quadro

        # Desenha o quadro branco
        c.setFillColor("white")  # Cor de fundo branca
        c.rect(pos_x_quadro, pos_y_quadro, largura_quadro, altura_quadro, fill=1)

        # Ajusta a posi√ß√£o do texto dentro do ret√¢ngulo
        pos_x_texto = pos_x_quadro + margem_interna
        pos_y_texto = pos_y_quadro + altura_quadro - margem_interna  # Alinhado no topo dentro do quadro

        # Adiciona a numera√ß√£o e a sigla dentro do quadro
        c.setFont(fonte, tamanho_fonte)
        c.setFillColor(red)  # Define a cor como vermelho
        c.drawString(pos_x_texto, pos_y_texto, texto_numero)

        c.setFont(fonte, tamanho_fonte - 2)
        c.drawString(pos_x_texto, pos_y_texto - 15, texto_sigla)
        c.drawString(pos_x_texto, pos_y_texto - 30, texto_matricula)

        c.save()
        buffer.seek(0)

        num_page = PdfReader(buffer).pages[0]
        pagina.merge_page(num_page)
        writer.add_page(pagina)

    # Salva o arquivo ajustado com numera√ß√£o
    with open(output_path, "wb") as f:
        writer.write(f)


def processar_pasta_docs(
        pasta_entrada="docs processados",
        pasta_saida="docs numerados",
        caminho_txt="sigla_para_a_numeracao.txt",
        inicio=2,
        fonte="Helvetica",
        tamanho_fonte=12
):
    """Processa todos os PDFs da pasta de entrada e os salva numerados na pasta de sa√≠da"""
    sigla, matricula = carregar_dados_txt(caminho_txt)

    # Garante que a pasta de sa√≠da exista
    os.makedirs(pasta_saida, exist_ok=True)

    for arquivo in os.listdir(pasta_entrada):
        if arquivo.endswith(".pdf"):
            input_path = os.path.join(pasta_entrada, arquivo)
            output_path = os.path.join(pasta_saida, arquivo)
            print(f"Numerando: {arquivo}")
            ajustar_orientacao_e_numerar(input_path, output_path, sigla, matricula, inicio, fonte, tamanho_fonte)


# Personalize os par√¢metros aqui
processar_pasta_docs(
    pasta_entrada="docs processados",
    pasta_saida="docs numerados",
    caminho_txt="sigla_para_a_numeracao.txt",  # Caminho para o arquivo .txt com a sigla e matr√≠cula
    inicio=2,  # In√≠cio da numera√ß√£o
    fonte="Helvetica-Bold",  # Fonte para o texto
    tamanho_fonte=10  # Tamanho da fonte principal
)


# ----------------------------------------------------------------
#C√≥digo 6 autuar os processos no GIAP;
# ----------------------------------------------------------------

# Fun√ß√£o para carregar credenciais de login
def carregar_credenciais(caminho_json):
    with open(caminho_json, 'r') as arquivo:
        return json.load(arquivo)

# Dicion√°rio de siglas
dicionario_orgao = {
    "TJSP": "TRIBUNAL DE JUSTI√áA DO ESTADO DE S√ÉO PAULO",
    "DEFENS. P√öBLICA de SP": "DEFENSORIA PUBLICA DO ESTADO DE SAO PAULO",
    "TRT 2": "TRIBUNAL REGIONAL DO TRABALHO DA 2A REGIAO",
    # Adicione outras siglas conforme necess√°rio
}

# Configura√ß√£o do WebDriver
def configurar_webdriver():
    options = webdriver.ChromeOptions()
    caminho_para_salvar_pdf = r"C:\Users\wesley\PycharmProjects\Autuar-processos\docs numerados"

    settings = {
        "recentDestinations": [{"id": "Save as PDF", "origin": "local", "account": ""}],
        "selectedDestinationId": "Save as PDF",
        "version": 2,
        "isHeaderFooterEnabled": False,
        "isCssBackgroundEnabled": True,
    }

    prefs = {
        "printing.print_preview_sticky_settings.appState": json.dumps(settings),
        "savefile.default_directory": caminho_para_salvar_pdf,
    }

    options.add_experimental_option("prefs", prefs)
    options.add_argument('--kiosk-printing')  # Imprime automaticamente

    return webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)


def renomear_todos_pdfs(diretorio, texto_busca, tentativas=5, intervalo=2):
    for arquivo in os.listdir(diretorio):
        if arquivo.endswith(".pdf"):
            caminho_pdf = os.path.join(diretorio, arquivo)

            # Ler o conte√∫do do PDF
            with fitz.open(caminho_pdf) as pdf:
                texto = "".join(pagina.get_text() for pagina in pdf)

            # Verificar se o texto buscado est√° no PDF
            if texto_busca in texto:
                inicio = texto.find(texto_busca) + len(texto_busca)
                fim = texto.find("\n", inicio)
                numero_processo = texto[inicio:fim].strip().replace("/", "-")
                novo_nome = f"capa_ref_pa_{numero_processo}.pdf"

                # Tentar renomear o arquivo
                for tentativa in range(tentativas):
                    try:
                        os.rename(caminho_pdf, os.path.join(diretorio, novo_nome))
                        print(f"Arquivo '{arquivo}' renomeado para: {novo_nome}")
                        break  # Sai do loop de tentativas ao renomear
                    except PermissionError as e:
                        print(f"Tentativa {tentativa + 1} falhou: {e}")
                        time.sleep(intervalo)
                else:
                    print(f"N√£o foi poss√≠vel renomear '{arquivo}' ap√≥s {tentativas} tentativas.")
            else:
                print(f"Texto n√£o encontrado no arquivo '{arquivo}'.")

# Fun√ß√£o para renomear todos os PDFs restantes
#def renomear_todos_pdfs(diretorio, texto_busca):
#    print("Renomeando todos os PDFs...")
#    for arquivo in os.listdir(diretorio):
#        if arquivo.endswith(".pdf") and not arquivo.startswith("capa_ref_pa_"):
#            renomear_capa_pdf(diretorio, texto_busca)
#    print("Renomea√ß√£o conclu√≠da.")



# Fun√ß√£o principal
def processar_planilha(caminho_excel, caminho_credenciais):
    workbook = openpyxl.load_workbook(caminho_excel)
    sheet = workbook.active

    credenciais = carregar_credenciais(caminho_credenciais)
    usuario = credenciais['usuario']
    senha = credenciais['senha']

    driver = configurar_webdriver()

    try:
        # Acessar o site e fazer login
        driver.get('https://carapicuiba.giap.com.br/apex/carapi/f?p=652:LOGIN')
        driver.find_element(By.ID, 'P101_USERNAME').send_keys(usuario)
        driver.find_element(By.ID, 'P101_PASSWORD').send_keys(senha)
        driver.find_element(By.ID, 'wwvFlowForm').submit()

        time.sleep(2)
        driver.find_element(By.XPATH, '//*[@id="report_R5001749296453489731"]/tbody/tr[2]/td/table/tbody/tr[2]/td[2]/a').click()

        # BTN de clicar na guia de processo
        driver.find_element(By.XPATH,
                            '// *[ @ id = "wwvFlowForm"] / div[2] / div / table / tbody / tr / td / div[1] / div[3] / div[1] / img').click()

        for row_index, row in enumerate(sheet.iter_rows(min_row=2), start=2):
            orgao = row[2].value
            informacoes = row[7].value

            if orgao is None or str(orgao).strip() == "":
                break

            significado = dicionario_orgao.get(orgao, "Sigla n√£o encontrada")

            driver.find_element(By.XPATH, '//*[@id="R5002551580972218898"]/tbody/tr[2]/td/ol/li[1]/a').click()
            time.sleep(1)
            # Preencher dados no GIAP

            campo_de_busca_nome_orgao = driver.find_element(By.XPATH, '//*[@id="P52_DSP_RESPONSAVEL"]')
            campo_de_busca_nome_orgao.clear()
            campo_de_busca_nome_orgao.send_keys(significado)
            time.sleep(1)
            campo_de_busca_nome_orgao.submit()
            time.sleep(3)

            # Selecionar op√ß√µes espec√≠ficas no GIAP (exemplo)
            if significado == "TRIBUNAL DE JUSTI√áA DO ESTADO DE S√ÉO PAULO":
                driver.find_element(By.XPATH,
                                    '//*[@id="report_P52_TIPO_EXPEDIENTE"]/tbody/tr[2]/td/table/tbody/tr[3]/td[10]/a').click()
            elif significado == "DEFENSORIA PUBLICA DO ESTADO DE SAO PAULO":
                driver.find_element(By.XPATH,
                                    '//*[@id="report_P52_TIPO_EXPEDIENTE"]/tbody/tr[2]/td/table/tbody/tr[2]/td[10]/a').click()
            elif significado == "TRIBUNAL REGIONAL DO TRABALHO DA 2A REGIAO":
                driver.find_element(By.XPATH,
                                    '//*[@id="report_P52_TIPO_EXPEDIENTE"]/tbody/tr[2]/td/table/tbody/tr[2]/td[10]/a').click()

            dropdown = driver.find_element(By.XPATH, '//*[@id="P51_TIPO_PROCESSO"]')
            select = Select(dropdown)
            select.select_by_visible_text("SEC. MUNICIPAL DE ASSUNTOS JUR√çDICOS")
            time.sleep(1)

            # Preencher informa√ß√µes no campo de texto
            driver.find_element(By.XPATH, '//*[@id="P51_COD_ASSUNTO1"]').send_keys('37')
            driver.find_element(By.XPATH, '//*[@id="P51_PRCS_DES_PROCESSO"]').send_keys(informacoes)
            time.sleep(1)

            #Clicar no Radio Button chamado de "Eletr√¥nico" e o selecionar com um click
            radio_button = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.XPATH, '//*[@id="P51_STA_ELETRO"]/div/div/div[1]'))
            )

            # Clica no radio button
            radio_button.click()


            # Clicar no bot√£o "gerar processo"
            driver.find_element(By.XPATH, '//*[@id="B5000650400896932596"]/span').click()

            # Copiar n√∫mero do PA gerado
            copiar_num_pa = WebDriverWait(driver, 10).until(
                EC.presence_of_element_located(
                    (By.XPATH, '//*[@id="apex_layout_5000650194595932588"]/tbody/tr/td[2]'))
            )
            valor_giap_copiado = copiar_num_pa.text

            # Atualizar a c√©lula correspondente na planilha
            sheet.cell(row=row_index, column=9, value=valor_giap_copiado)  # Coluna 9 √© "N√∫mero PA gerado"

            # Salvar a planilha atualizada
            workbook.save(caminho_excel)

            # Daqui para baixo segue o meu c√≥digo para baixar a capa como PDF

            # Salvar a p√°gina como PDF
            time.sleep(2)
            driver.find_element(By.XPATH, '//*[@id="B4985322504071026678"]/span').click()

            # Identificar as janelas/abas abertas
            janelas = driver.window_handles
            driver.switch_to.window(janelas[1])
            time.sleep(2)
            print("Salvando a p√°gina como PDF...")
            driver.execute_script('window.print();')
            time.sleep(2)
            driver.close()

            # Obter todos os identificadores de janela
            print("Capa do processo salva como PDF com sucesso!")

            time.sleep(2)

            # Voltar para a janela principal
            driver.switch_to.window(janelas[0])
            time.sleep(2)

            # Clicar na engrenagem para aparecer o menu para voltar para a tela principal do GIAP
            driver.find_element(By.XPATH, '//*[@id="menu_app"]').click()
            time.sleep(2)
            # Depois de clicar na engrenagem, vai clicar no bot√£o chamado de "Processo" para voltar para a tela principal do GIAP
            driver.find_element(By.XPATH,
                                '//*[@id="aparece_app"]/div[2]/table/tbody/tr/td/div[2]/div/div/div[1]/div[1]/img').click()
            time.sleep(2)

        driver.quit()
        renomear_todos_pdfs(r"C:\Users\wesley\PycharmProjects\Autuar-processos\docs numerados",
                          "INFORMA√á√ïES DO PROCESSO - ")


    except Exception as e:
        print(f"Erro ocorreu: {e}")
    finally:
        print("Processos autuados e capas imprimidas e renomeadas.")
        #driver.quit()

# Caminhos
caminho_excel = "dados_para_autuar_processos.xlsx"
caminho_credenciais = "credenciais_login_GIAP.json"

# Executar
processar_planilha(caminho_excel, caminho_credenciais)



# ----------------------------------------------------------------
#C√≥digo 6.1 inserir uma pg em branco nas capas do GIAP
#Isso √© por causa da impressora que imprime frente e verso
# ----------------------------------------------------------------






def adicionar_pagina_branca_em_pdfs(diretorio):
    """
    Insere uma p√°gina em branco como p√°gina 2 em PDFs cujo nome come√ßa com 'capa_ref_pa_'.

    :param diretorio: Caminho do diret√≥rio onde est√£o os arquivos PDF.
    """
    # Caminho do diret√≥rio com os PDFs
    for arquivo in os.listdir(diretorio):
        if arquivo.startswith("capa_ref_pa_") and arquivo.endswith(".pdf"):
            caminho_arquivo = os.path.join(diretorio, arquivo)

            # Ler o PDF existente
            reader = PdfReader(caminho_arquivo)
            writer = PdfWriter()

            # Adicionar a primeira p√°gina do PDF original
            writer.add_page(reader.pages[0])

            # Adicionar uma p√°gina em branco
            writer.add_blank_page()

            # Adicionar as p√°ginas restantes do PDF original
            for pagina in reader.pages[1:]:
                writer.add_page(pagina)

            # Salvar o novo arquivo no mesmo local com o mesmo nome
            with open(caminho_arquivo, "wb") as arquivo_pdf:
                writer.write(arquivo_pdf)

            print(f"P√°gina em branco adicionada ao arquivo: {arquivo}")

# Caminho do diret√≥rio
diretorio_pdfs = r"C:\Users\wesley\PycharmProjects\Autuar-processos\docs numerados"

# Executar a fun√ß√£o
adicionar_pagina_branca_em_pdfs(diretorio_pdfs)



# ----------------------------------------------------------------
#C√≥digo 7 - Juntar capa com os documentos em si
# ----------------------------------------------------------------






# Fun√ß√£o para verificar se um n√∫mero segue o padr√£o especificado
def verificar_padrao_numero(numero):
    padrao = r"\d{7}-\d{2}\.\d{4}\.\d\.\d{2}\.\d{4}"
    return bool(re.match(padrao, numero))


# Fun√ß√£o para processar o diret√≥rio e os arquivos do Excel
def processar_documentos(diretorio_docs, arquivo_excel):
    # Carregar o arquivo Excel e a guia 'CAPTACOES'
    wb = load_workbook(arquivo_excel)
    ws = wb["CAPTACOES"]

    # Iterar sobre as linhas da planilha (ignorando o cabe√ßalho)
    for row in ws.iter_rows(min_row=2, values_only=True):
        if all(cell is None for cell in row):
            break  # Parar se encontrar uma linha completamente vazia

        orgao = row[2]  # Coluna "Org√£o"
        id_processo = row[0]  # Coluna "ID processo"
        numero_pa_gerado = row[8]  # Coluna "Numero PA gerado"

        # Verificar as condi√ß√µes da linha
        if orgao == "TJSP" and id_processo and verificar_padrao_numero(id_processo):
            # Tratar o valor de 'numero_pa_gerado'
            numero_processo = numero_pa_gerado.replace("/", "-")
            nome_capa = f"capa_ref_pa_{numero_processo}.pdf"
            caminho_capa = os.path.join(diretorio_docs, nome_capa)
            caminho_documento = os.path.join(diretorio_docs, f"{id_processo}.pdf")

            # Verificar se a capa e o documento existem
            if os.path.exists(caminho_capa) and os.path.exists(caminho_documento):
                # Criar o nome do arquivo combinado
                nome_saida = f"{id_processo}_completo.pdf"
                caminho_saida = os.path.join(diretorio_docs, nome_saida)

                # Combinar os PDFs
                merger = PdfMerger()
                merger.append(caminho_capa)
                merger.append(caminho_documento)
                merger.write(caminho_saida)
                merger.close()

                # Mover a capa processada para a lixeira
                send2trash(caminho_capa)

                # Mover o documento principal para a lixeira
                send2trash(caminho_documento)

                print(f"Processado: {nome_saida}")
            else:
                print(f"Arquivos ausentes para o ID {id_processo}.")
        elif id_processo and id_processo.startswith("/// "):
            numero_potencial = id_processo[4:].strip()
            if verificar_padrao_numero(numero_potencial):
                # Repetir o mesmo processo para as linhas com "/// "
                numero_processo = numero_pa_gerado.replace("/", "-")
                nome_capa = f"capa_ref_pa_{numero_processo}.pdf"
                caminho_capa = os.path.join(diretorio_docs, nome_capa)
                caminho_documento = os.path.join(diretorio_docs, f"{numero_potencial}.pdf")

                # Verificar se a capa e o documento existem
                if os.path.exists(caminho_capa) and os.path.exists(caminho_documento):
                    nome_saida = f"{numero_potencial}_completo.pdf"
                    caminho_saida = os.path.join(diretorio_docs, nome_saida)

                    # Combinar os PDFs
                    merger = PdfMerger()
                    merger.append(caminho_capa)
                    merger.append(caminho_documento)
                    merger.write(caminho_saida)
                    merger.close()

                    # Mover a capa processada para a lixeira
                    send2trash(caminho_capa)

                    # Mover o documento principal para a lixeira
                    send2trash(caminho_documento)

                    print(f"Processado: {nome_saida}")
                else:
                    print(f"Arquivos ausentes para o n√∫mero potencial {numero_potencial}.")


# Definir os caminhos
diretorio_docs = "docs numerados"
arquivo_excel = "dados_para_autuar_processos.xlsx"

# Executar o processamento
processar_documentos(diretorio_docs, arquivo_excel)


# ----------------------------------------------------------------
#C√≥digo 8 Criar o modelo de registro para os cadernos SAJ de procurador
# ----------------------------------------------------------------






# Caminhos para os arquivos
excel_file = 'dados_para_autuar_processos.xlsx'
sheet_name = 'CAPTACOES'
word_template = 'modelos_pa_procuradoria.docx'
output_folder = Path("docs numerados")
output_folder.mkdir(exist_ok=True)  # Garante que a pasta existe

try:
    # Carregar o Excel
    df = pd.read_excel(excel_file, sheet_name=sheet_name)

    # Obter a data atual
    current_day = datetime.today().day
    current_month = datetime.today().month
    months = ['Jan', 'Fev', 'Mar', 'Abr', 'Mai', 'Jun', 'Jul', 'Ago', 'Set', 'Out', 'Nov', 'Dez']
    current_month_text = months[current_month - 1]
    current_year = datetime.today().year
    current_date = f"{current_day} {current_month_text} {current_year}"

    # Filtrar apenas as linhas que t√™m um procurador definido
    df_filtered = df[df['Procurador'].notna()]

    # Criar lista de registros para inserir no Word
    table_data = []
    for _, row in df_filtered.iterrows():
        data_procurador = f"{current_date} - {row['Procurador']}"
        table_data.append({
            'N√∫mero PA Gerado': row['Numero PA gerado'],
            'Informa√ß√µes Completas para Autuar o PA': row['Informacoes completas para autuar o PA'],
            'Data Atual + Procurador': data_procurador,
            'Livro saj': row['Livro saj']
        })

    # Vari√°vel para numerar os arquivos
    file_count = 1

    # Criar arquivos separados
    for i in range(0, len(table_data), 4):  # Processa at√© 3 registros por arquivo
        doc = DocxTemplate(word_template)

        # Pegamos um lote de 10 registros (ou menos, se for o final da lista)
        context = {'table': table_data[i:i+4]}

        # Renderizar o documento com os dados
        doc.render(context)

        # Criar um novo documento Word a partir do preenchido
        final_doc = Document()
        final_doc.add_paragraph().add_run().add_break()  # Adiciona quebra de p√°gina
        for element in doc.element.body:
            final_doc.element.body.append(element)

        # Adicionar uma p√°gina em branco
        final_doc.add_page_break()

        # Nome do arquivo com numera√ß√£o
        output_file = output_folder / f"modelo_registros_PA_procuradoria_preenchido{file_count}.docx"

        # Salvar o documento final
        final_doc.save(output_file)
        print(f"‚úÖ Documento gerado com sucesso: {output_file}")

        file_count += 1  # Incrementa o n√∫mero do arquivo

except Exception as e:
    print(f"‚ùå Erro ao gerar o documento: {e}")




# ----------------------------------------------------------------
#C√≥digo 8.1 impressora 15/04/2025
# ----------------------------------------------------------------

#Como obter o nome da impressora? 1 - Apertar tecla windows e escrever a palavra "impressora", logo aparecer√° a guia "impressoras e scanners", clique.
# 2 - Escolha a impressora a ser usada para realizar as impress√µes nela, d√™ 1 clique com o bot√£o esquerdo e clique na op√ß√£o "Gerenciar"
# 3 - Na tela que abriu, clique na op√ß√£o "Propriedades da impressora"
# 4 - Na tela que abrir, ter√° o nome da impressora selecionado, copie e cole entre as aspas duplas NOME_IMPRESSORA_DESEJADA = "Lexmark MX720 Series XL"


# Nome fixo da impressora desejada
NOME_IMPRESSORA_DESEJADA = "Lexmark MX720 Series XL"

# Diret√≥rio dos arquivos a serem impressos
caminho = r"C:\Users\wesley\PycharmProjects\Autuar-processos\docs numerados"

print("Abaixo segue a listagem das impressoras dispon√≠veis:")

# Lista todas as impressoras dispon√≠veis
lista_impressoras = win32print.EnumPrinters(2)
print("---------------------")
for i, impressora in enumerate(lista_impressoras):
    print(f"{i}: {impressora[2]}")  # O nome da impressora est√° na posi√ß√£o [2]
print("---------------------")

# Seleciona a impressora pelo nome
impressora = None
for imp in lista_impressoras:
    if NOME_IMPRESSORA_DESEJADA.lower() in imp[2].lower():
        impressora = imp
        break

if impressora:
    win32print.SetDefaultPrinter(impressora[2])
    print(f"Impressora selecionada: {impressora[2]}")


    # Fun√ß√£o para verificar se h√° documentos na fila de impress√£o
    def obter_numero_jobs():
        try:
            printer_info = win32print.OpenPrinter(impressora[2])
            jobs = win32print.EnumJobs(printer_info, 0, -1, 1)  # Obt√©m a lista de trabalhos na fila
            win32print.ClosePrinter(printer_info)
            return len(jobs)  # Retorna o n√∫mero de trabalhos na fila
        except Exception as e:
            print(f"Erro ao acessar a fila de impress√£o: {e}")
            return -1


    # Fun√ß√£o para verificar se um arquivo est√° desbloqueado
    def arquivo_desbloqueado(caminho_arquivo, timeout=30):
        tempo_inicio = time.time()
        while time.time() - tempo_inicio < timeout:
            try:
                with open(caminho_arquivo, "r"):
                    return True
            except PermissionError:
                print(f"Aguardando o arquivo {os.path.basename(caminho_arquivo)} ser liberado...")
                time.sleep(5)
        print(f"Aviso: Tempo limite atingido para desbloqueio de {os.path.basename(caminho_arquivo)}.")
        return False


    # Lista inicial de arquivos
    lista_arquivos = sorted(os.listdir(caminho))  # Ordena os arquivos por nome
    arquivos_nao_impressos = []

    print("Arquivos encontrados para impress√£o:")
    for arquivo in lista_arquivos:
        print(arquivo)

    while lista_arquivos:
        arquivo = lista_arquivos[0]  # Pega o primeiro arquivo da lista
        caminho_arquivo = os.path.join(caminho, arquivo)

        try:
            print(f"Enviando para impress√£o: {arquivo}")
            win32api.ShellExecute(0, "print", caminho_arquivo, None, caminho, 0)
        except Exception as e:
            print(f"Erro ao tentar imprimir {arquivo}: {e}")
            arquivos_nao_impressos.append(arquivo)
            lista_arquivos.pop(0)
            continue  # Passa para o pr√≥ximo arquivo

        # Esperar at√© que o trabalho entre na fila de impress√£o (com timeout)
        print("Aguardando o arquivo entrar na fila de impress√£o...")
        tempo_inicio = time.time()
        while obter_numero_jobs() == 0:
            if time.time() - tempo_inicio > 30:
                print(f"Aviso: {arquivo} n√£o entrou na fila de impress√£o a tempo.")
                arquivos_nao_impressos.append(arquivo)
                break
            time.sleep(1)

        # Aguardando a impress√£o ser conclu√≠da (com timeout)
        print("Aguardando a impress√£o ser conclu√≠da...")
        tempo_inicio = time.time()
        while obter_numero_jobs() > 0:
            if time.time() - tempo_inicio > 120:
                print(f"Aviso: Tempo limite atingido para impress√£o de {arquivo}.")
                arquivos_nao_impressos.append(arquivo)
                break
            time.sleep(5)

        print(f"Verificando se o arquivo {arquivo} est√° desbloqueado...")
        if not arquivo_desbloqueado(caminho_arquivo):
            arquivos_nao_impressos.append(arquivo)
            lista_arquivos.pop(0)
            continue

        print(f"Movendo {arquivo} para a lixeira...")
        send2trash(caminho_arquivo)
        lista_arquivos = sorted(os.listdir(caminho))  # Atualiza a lista de arquivos restantes


    # Exibir arquivos que n√£o foram impressos
    if arquivos_nao_impressos:
        print("Os seguintes arquivos n√£o foram impressos e precisam ser verificados manualmente:")
        for arquivo in arquivos_nao_impressos:
            print(f"- {arquivo}")
else:
    print("A impressora descrita na variavel N√ÉO FOI localizada, voc√™ dever√° imprimir manualmente DEPOIS que o c√≥digo terminar de executar.")




# ----------------------------------------------------------------
#C√≥digo 9 atualizar as planilhas de judicial e do expediente
# ----------------------------------------------------------------




# Fun√ß√£o para autenticar usando uma conta de servi√ßo
def autenticar_google_sheets():
    # Caminho para o arquivo JSON das credenciais
    credenciais_json = 'chave.json'

    # Escopos necess√°rios para acessar Google Sheets e Drive
    escopos = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']

    # Configura√ß√£o das credenciais
    credenciais = ServiceAccountCredentials.from_json_keyfile_name(credenciais_json, escopos)
    cliente = gspread.authorize(credenciais)
    return cliente


# Fun√ß√£o para encontrar a pr√≥xima linha vazia
def encontrar_proxima_linha_vazia(planilha):
    valores = planilha.col_values(2)  # Coluna "Processo Administrativo" como refer√™ncia
    return len(valores) + 1


# Fun√ß√£o para processar o Excel e atualizar Google Sheets
def processar_excel_para_google_planilhas():
    # Caminho do arquivo Excel
    caminho_excel = 'dados_para_autuar_processos.xlsx'

    # Abrindo o arquivo Excel
    workbook = openpyxl.load_workbook(caminho_excel)
    sheet = workbook.active

    # Autenticando e conectando ao Google Sheets
    cliente = autenticar_google_sheets()

    # Acessando a planilha Google para o Caso 1
    planilha_caso_1 = cliente.open_by_key('1B-zoQMhTcQM4mSSKeQdHJj6uidH_q_NfkvkJYrrTMoU').worksheet(
        'Processos Judiciais')

    # Processando as linhas do Excel
    for row in sheet.iter_rows(min_row=2, values_only=True):
        # Verificar se a linha est√° completamente vazia
        if all(cell is None for cell in row):
            break

        # Dados das colunas
        id_processo = row[0]  # Primeira coluna
        procurador = row[1]  # Segunda coluna
        livro_saj = row[5]  # Sexta coluna
        info_autuar_pa = row[7]  # Oitava coluna
        numero_pa_gerado = row[8]  # Nona coluna

        # Verificar se o ID do processo come√ßa com '///'
        if isinstance(id_processo, str) and id_processo.startswith('///'):
            # Extrair o conte√∫do entre as barras
            id_processo = id_processo.split('/')[3]

        # Caso 1: Se "Procurador" estiver preenchido
        if procurador:
            data_atual = datetime.now().strftime('%d/%m/%Y')
            linha_caso_1 = [
                numero_pa_gerado,  # Coluna "PA"
                id_processo,  # Coluna "N√öMERO DE PROCESSO JUDICIAL"
                info_autuar_pa,  # Coluna "ASSUNTO"
                f"{procurador} {data_atual}",  # Coluna "andamento"
                procurador,  # Coluna "PROCURADOR RESPONS√ÅVEL"
                livro_saj  # Coluna "LIVRO SAJ"
            ]
            planilha_caso_1.append_row(linha_caso_1)

        # Caso 2: Se "Procurador" estiver vazio, pula para a pr√≥xima linha
        else:
            continue

    print("Processos inseridos na planilha da procuradoria com sucesso!")


# Fun√ß√£o para processar o Caso 2
def processar_caso_2():
    # Caminho do arquivo Excel
    caminho_excel = 'dados_para_autuar_processos.xlsx'

    # Abrindo o arquivo Excel
    workbook = openpyxl.load_workbook(caminho_excel)
    sheet = workbook.active

    # Autenticando e conectando ao Google Sheets
    cliente = autenticar_google_sheets()

    # Acessando a planilha Google do Caso 2
    planilha_caso_2 = cliente.open_by_key('16_zlC5bRdyGTqFcVFvRIBCYzP-fjoPN9i64tD5DGe5c').worksheet('Atual')

    # Processando as linhas do Excel
    for row in sheet.iter_rows(min_row=2, values_only=True):
        # Verificar se a linha est√° completamente vazia
        if all(cell is None for cell in row):
            break

        # Dados das colunas
        id_processo = row[0]  # Primeira coluna
        procurador = row[1]  # Segunda coluna
        numero_pa_gerado = row[8]  # Nona coluna
        orgao = row[2]  # Terceira coluna (para "√ìrg√£o de Destino")

        # Verificar se o ID do processo come√ßa com '///'
        if isinstance(id_processo, str) and id_processo.startswith('///'):
            # Extrair o conte√∫do entre as barras
            id_processo = id_processo.split('/')[3]

        # Pular a linha se "Procurador" estiver preenchido
        if procurador:
            continue

        # Encontrar a pr√≥xima linha vazia
        proxima_linha = encontrar_proxima_linha_vazia(planilha_caso_2)

        # Atualizar as colunas manualmente
        planilha_caso_2.update_cell(proxima_linha, 2, numero_pa_gerado)  # Coluna "Processo Administrativo"
        planilha_caso_2.update_cell(proxima_linha, 3, id_processo)  # Coluna "Refer√™ncia"
        planilha_caso_2.update_cell(proxima_linha, 4, orgao)  # Coluna "√ìrg√£o de Destino"
        planilha_caso_2.update_cell(proxima_linha, 14, "Aguardando provid√™ncias")
        # Atualizar a c√©lula com a mensagem e a data
        planilha_caso_2.update_cell(proxima_linha, 15, f"PA autuado em: {datetime.now().strftime('%d/%m/%Y')}")

    print("Processos atualizados com sucesso na planilha de controle de Of√≠cios")


# Executar os processos
processar_excel_para_google_planilhas()
processar_caso_2()



# ----------------------------------------------------------------
#C√≥digo 10 - Limpar os itens da planilha excel e diret√≥rios para a pr√≥xima execu√ß√£o
# ----------------------------------------------------------------





def limpar_planilha_captacoes():
    def contagem_regressiva():
        for t in range(30, 0, -1):
            print(f"\rAs informa√ß√µes da aba 'CAPTACOES' do arquivo Excel ser√£o apagadas em {t} segundos...", end="")
            time.sleep(1)
        print("\nTempo esgotado. Limpando a aba automaticamente para futuros usos.")
        realizar_limpeza()

    def realizar_limpeza():
        # Caminho do arquivo Excel
        caminho_arquivo = "dados_para_autuar_processos.xlsx"

        # Abrindo o arquivo e selecionando a aba "CAPTACOES"
        workbook = load_workbook(caminho_arquivo)
        planilha = workbook["CAPTACOES"]

        # Iterando sobre as linhas, preservando o cabe√ßalho
        for row in planilha.iter_rows(min_row=2, min_col=1, max_col=9):  # Colunas A (1) at√© I (9)
            for cell in row:
                cell.value = None  # Apaga o conte√∫do da c√©lula

        # Salvando as modifica√ß√µes
        workbook.save(caminho_arquivo)
        print("As informa√ß√µes da aba 'CAPTACOES' foram limpas para a pr√≥xima execu√ß√£o.")

    # Inicia a contagem regressiva em uma nova thread
    thread_timer = threading.Thread(target=contagem_regressiva)
    thread_timer.start()


# Chama a fun√ß√£o para execu√ß√£o
limpar_planilha_captacoes()




# ----------------------------------------------------------------
#C√≥digo 10.1 - mover  o conteudo dos diret√≥rios para a lixeira
# ----------------------------------------------------------------



def mover_para_lixeira():
    # Lista das pastas que ter√£o seus conte√∫dos movidos para a lixeira
    pastas = ["docs PJs", "docs numerados", "docs processados"]

    for pasta in pastas:
        # Verifica se a pasta existe
        if not os.path.exists(pasta):
            print(f"A pasta '{pasta}' n√£o foi encontrada.")
            continue

        # Itera sobre os arquivos e subdiret√≥rios na pasta
        for item in os.listdir(pasta):
            caminho_item = os.path.join(pasta, item)

            # Move para a lixeira apenas arquivos ou diret√≥rios v√°lidos
            if os.path.isfile(caminho_item) or os.path.isdir(caminho_item):
                tentativas = 0
                max_tentativas = 20  # N√∫mero m√°ximo de tentativas
                while tentativas < max_tentativas:
                    try:
                        # Verifica se o arquivo est√° em uso tentando abrir em modo exclusivo
                        with open(caminho_item, 'rb+'):
                            pass
                        # Se o arquivo n√£o estiver em uso, move para a lixeira
                        send2trash(caminho_item)
                        print(f"Movido para a lixeira: {caminho_item}")
                        break
                    except PermissionError:
                        tentativas += 1
                        print(f"O arquivo '{caminho_item}' est√° em uso. Tentando novamente ({tentativas}/{max_tentativas})...")
                        time.sleep(5)  # Aguarda 5 segundos antes de tentar novamente
                    except Exception as e:
                        print(f"Erro ao tentar mover '{caminho_item}' para a lixeira: {e}")
                        break
                else:
                    print(f"O arquivo '{caminho_item}' n√£o p√¥de ser movido ap√≥s {max_tentativas} tentativas.")

    print("Processo de mover arquivos para a lixeira conclu√≠do.")

# Chamar a fun√ß√£o
mover_para_lixeira()
